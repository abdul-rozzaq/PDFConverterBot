import zipfile
from typing import List
from pathlib import Path
import asyncio
import aiofiles
from PIL import Image, ImageOps
import fitz  # PyMuPDF
from docx import Document
import pytesseract
import easyocr
from pdf2image import convert_from_path
import openpyxl
from app.core.config import settings


class DocumentConverter:
    def __init__(self):
        self.temp_dir = Path(settings.TEMP_DIR)
        self.temp_dir.mkdir(exist_ok=True)

    async def pdf_to_word(self, pdf_path: str) -> str:
        """Convert PDF to Word document"""
        try:
            # Open PDF
            doc = fitz.open(pdf_path)

            # Create Word document
            word_doc = Document()

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()

                if text.strip():
                    word_doc.add_paragraph(text)

                # Add page break except for last page
                if page_num < len(doc) - 1:
                    word_doc.add_page_break()

            # Save Word document
            output_path = pdf_path.replace(".pdf", ".docx")
            word_doc.save(output_path)
            doc.close()

            return output_path

        except Exception as e:
            raise Exception(f"PDF to Word conversion failed: {str(e)}")

    async def word_to_pdf(self, docx_path: str) -> str:
        """Convert Word document to PDF"""
        try:
            # This is a simplified version. For production, consider using LibreOffice
            # or Microsoft Word API for better formatting preservation

            doc = Document(docx_path)
            text_content = []

            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)

            # Create PDF using reportlab or similar library
            # For now, we'll create a simple text-based PDF using fitz
            pdf_doc = fitz.open()
            page = pdf_doc.new_page()

            full_text = "\n".join(text_content)
            text_rect = fitz.Rect(72, 72, 540, 720)  # Margins
            page.insert_textbox(text_rect, full_text, fontsize=12)

            output_path = docx_path.replace(".docx", ".pdf")
            pdf_doc.save(output_path)
            pdf_doc.close()

            return output_path

        except Exception as e:
            raise Exception(f"Word to PDF conversion failed: {str(e)}")

    async def pdf_to_images(self, pdf_path: str, format: str = "PNG") -> List[str]:
        """Convert PDF pages to images"""
        try:
            # Convert PDF to images
            images = convert_from_path(pdf_path, dpi=300)

            image_paths = []
            base_name = Path(pdf_path).stem

            for i, image in enumerate(images):
                image_path = self.temp_dir / f"{base_name}_page_{i+1}.{format.lower()}"
                image.save(image_path, format)
                image_paths.append(str(image_path))

            return image_paths

        except Exception as e:
            raise Exception(f"PDF to images conversion failed: {str(e)}")

    async def images_to_pdf(self, image_paths: List[str]) -> str:
        """Convert multiple images to PDF"""
        try:
            if not image_paths:
                raise ValueError("No images provided")

            images = []
            for img_path in image_paths:
                img = Image.open(img_path)
                # Convert to RGB if necessary
                if img.mode != "RGB":
                    img = img.convert("RGB")
                images.append(img)

            # Save as PDF
            output_path = self.temp_dir / "converted_images.pdf"
            images[0].save(output_path, save_all=True, append_images=images[1:])

            return str(output_path)

        except Exception as e:
            raise Exception(f"Images to PDF conversion failed: {str(e)}")

    async def pdf_to_text(self, pdf_path: str) -> str:
        """Extract text from PDF"""
        try:
            doc = fitz.open(pdf_path)
            text_content = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                text_content.append(f"--- Page {page_num + 1} ---\n{text}\n")

            doc.close()

            # Save as text file
            output_path = pdf_path.replace(".pdf", ".txt")
            async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
                await f.write("\n".join(text_content))

            return output_path

        except Exception as e:
            raise Exception(f"PDF to text conversion failed: {str(e)}")

    async def excel_to_pdf(self, excel_path: str) -> str:
        """Convert Excel to PDF"""
        try:
            # Read Excel file
            wb = openpyxl.load_workbook(excel_path)

            # Create PDF document
            pdf_doc = fitz.open()

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]

                # Create new page for each sheet
                page = pdf_doc.new_page()

                # Convert sheet data to text
                sheet_text = f"Sheet: {sheet_name}\n\n"
                for row in ws.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    sheet_text += row_text + "\n"

                # Insert text into PDF
                text_rect = fitz.Rect(72, 72, 540, 720)
                page.insert_textbox(text_rect, sheet_text, fontsize=10)

            output_path = excel_path.replace(".xlsx", ".pdf")
            pdf_doc.save(output_path)
            pdf_doc.close()

            return output_path

        except Exception as e:
            raise Exception(f"Excel to PDF conversion failed: {str(e)}")


class ImageProcessor:
    def __init__(self):
        self.temp_dir = Path(settings.TEMP_DIR)
        self.temp_dir.mkdir(exist_ok=True)

    async def compress_image(self, image_path: str, quality: int = 85) -> str:
        """Compress image"""
        try:
            with Image.open(image_path) as img:
                # Optimize and compress
                output_path = self.temp_dir / f"compressed_{Path(image_path).name}"

                if img.format == "PNG":
                    img.save(output_path, "PNG", optimize=True)
                else:
                    img.save(output_path, "JPEG", quality=quality, optimize=True)

                return str(output_path)

        except Exception as e:
            raise Exception(f"Image compression failed: {str(e)}")

    async def convert_to_grayscale(self, image_path: str) -> str:
        """Convert image to grayscale"""
        try:
            with Image.open(image_path) as img:
                grayscale_img = ImageOps.grayscale(img)

                output_path = self.temp_dir / f"grayscale_{Path(image_path).name}"
                grayscale_img.save(output_path)

                return str(output_path)

        except Exception as e:
            raise Exception(f"Grayscale conversion failed: {str(e)}")

    async def resize_image(self, image_path: str, width: int, height: int) -> str:
        """Resize image"""
        try:
            with Image.open(image_path) as img:
                resized_img = img.resize((width, height), Image.Resampling.LANCZOS)

                output_path = self.temp_dir / f"resized_{width}x{height}_{Path(image_path).name}"
                resized_img.save(output_path)

                return str(output_path)

        except Exception as e:
            raise Exception(f"Image resize failed: {str(e)}")

    async def convert_format(self, image_path: str, target_format: str) -> str:
        """Convert image format"""
        try:
            with Image.open(image_path) as img:
                # Convert to RGB if saving as JPEG
                if target_format.upper() == "JPEG" and img.mode in ("RGBA", "LA", "P"):
                    img = img.convert("RGB")

                output_path = self.temp_dir / f"{Path(image_path).stem}.{target_format.lower()}"
                img.save(output_path, target_format.upper())

                return str(output_path)

        except Exception as e:
            raise Exception(f"Format conversion failed: {str(e)}")


class OCRProcessor:
    def __init__(self):
        self.temp_dir = Path(settings.TEMP_DIR)
        self.temp_dir.mkdir(exist_ok=True)
        # Initialize EasyOCR reader
        self.reader = easyocr.Reader(["en", "ru"])  # Add more languages as needed

    async def extract_text_tesseract(self, image_path: str) -> str:
        """Extract text using Tesseract OCR"""
        try:
            # Configure Tesseract for better accuracy
            custom_config = r"--oem 3 --psm 6"
            text = pytesseract.image_to_string(Image.open(image_path), config=custom_config)

            # Save extracted text
            output_path = self.temp_dir / f"ocr_text_{Path(image_path).stem}.txt"
            async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
                await f.write(text)

            return str(output_path)

        except Exception as e:
            raise Exception(f"OCR extraction failed: {str(e)}")

    async def extract_text_easyocr(self, image_path: str) -> str:
        """Extract text using EasyOCR"""
        try:
            # Run OCR in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            results = await loop.run_in_executor(None, self.reader.readtext, image_path)

            # Extract text from results
            extracted_text = "\n".join([result[1] for result in results])

            # Save extracted text
            output_path = self.temp_dir / f"easyocr_text_{Path(image_path).stem}.txt"
            async with aiofiles.open(output_path, "w", encoding="utf-8") as f:
                await f.write(extracted_text)

            return str(output_path)

        except Exception as e:
            raise Exception(f"EasyOCR extraction failed: {str(e)}")


async def create_zip_archive(file_paths: List[str], archive_name: str) -> str:
    """Create ZIP archive from multiple files"""
    try:
        temp_dir = Path(settings.TEMP_DIR)
        zip_path = temp_dir / f"{archive_name}.zip"

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            for file_path in file_paths:
                zipf.write(file_path, Path(file_path).name)

        return str(zip_path)

    except Exception as e:
        raise Exception(f"ZIP creation failed: {str(e)}")
